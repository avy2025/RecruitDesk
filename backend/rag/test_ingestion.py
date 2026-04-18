import os
import sys
import unittest
import numpy as np
import json
import shutil
import tempfile
from typing import List, Dict, Any

# Add the backend directory to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from rag.schema import RAGDocument, DocumentMetadata
from rag.embedder import RAGEmbedder
from rag.vector_store import RAGVectorStore
from rag.loader import ResumeLoader
from rag.parser import parse_file

class TestRAGIngestion(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.embedder = RAGEmbedder.get_instance()
        cls.test_dir = tempfile.mkdtemp()
        cls.index_path = os.path.join(cls.test_dir, "test_index.faiss")
        cls.metadata_path = os.path.join(cls.test_dir, "test_metadata.json")

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.test_dir)

    def setUp(self):
        self.vector_store = RAGVectorStore(dimension=384)
        self.loader = ResumeLoader(self.vector_store, self.embedder)

    def create_fake_docx(self, path, text):
        import docx
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(path)

    def test_docx_ingestion(self):
        """Test ingesting a DOCX file."""
        docx_path = os.path.join(self.test_dir, "sample.docx")
        text = "Jane Doe. Senior Python Developer with 8 years of experience in AI and Cloud."
        self.create_fake_docx(docx_path, text)
        
        result = self.loader.ingest(docx_path, candidate_id="cand_jane_001")
        
        self.assertEqual(result["status"], "success")
        self.assertEqual(result["candidate_id"], "cand_jane_001")
        self.assertGreater(len(result["skills_detected"]), 0)
        self.assertIn("python", [s.lower() for s in result["skills_detected"]])
        
        # Verify vector store
        self.assertEqual(self.vector_store.index.ntotal, 1) # Should be 1 chunk for small text
        
    def test_small_file_single_chunk(self):
        """Test that very small files are stored as a single chunk."""
        docx_path = os.path.join(self.test_dir, "small.docx")
        text = "Short resume."
        self.create_fake_docx(docx_path, text)
        
        result = self.loader.ingest(docx_path)
        self.assertEqual(result["chunks_created"], 1)

    def test_duplicate_ingestion(self):
        """Test ingesting the same candidate twice."""
        docx_path = os.path.join(self.test_dir, "dup.docx")
        text = "Duplicate Candidate info."
        self.create_fake_docx(docx_path, text)
        
        # First ingestion
        self.loader.ingest(docx_path, candidate_id="dup_001")
        count1 = self.vector_store.index.ntotal
        
        # Second ingestion (same ID)
        # Note: Current implementation doesn't prevent duplicates yet, 
        # but we check if it correctly adds them as new entries if desired or fails.
        # Requirement was just "Verify behavior".
        self.loader.ingest(docx_path, candidate_id="dup_001")
        count2 = self.vector_store.index.ntotal
        
        self.assertEqual(count2, count1 + 1) # Currently it just adds more chunks

    def test_persistence_integrity(self):
        """Test index/metadata count matching after save and reload."""
        docx_path = os.path.join(self.test_dir, "persist.docx")
        text = "Persistence test candidate."
        self.create_fake_docx(docx_path, text)
        
        self.loader.ingest(docx_path)
        
        # Save
        idx_p = os.path.join(self.test_dir, "p.faiss")
        meta_p = os.path.join(self.test_dir, "p.json")
        self.vector_store.save(idx_p, meta_p)
        
        # Reload
        new_vs = RAGVectorStore(dimension=384)
        new_vs.load(idx_p, meta_p)
        
        integrity = new_vs.validate_integrity()
        self.assertTrue(integrity["is_consistent"])
        self.assertEqual(integrity["index_count"], self.vector_store.index.ntotal)
        
        # Verify search works on reloaded store
        query_emb = self.embedder.embed_text("Persistence test")
        results = new_vs.search(query_emb, k=1)
        self.assertGreater(len(results), 0)
        self.assertIn("Persistence", results[0][0]["text"])

    def test_empty_file_rejection(self):
        """Test that empty or unparseable files are rejected."""
        empty_path = os.path.join(self.test_dir, "empty.docx")
        # Create an almost empty docx
        import docx
        doc = docx.Document()
        doc.save(empty_path)
        
        result = self.loader.ingest(empty_path)
        self.assertEqual(result["status"], "error")
        self.assertIn("empty", result["message"].lower())

if __name__ == "__main__":
    unittest.main()
